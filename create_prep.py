from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_prep_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('Quiz & Viva Preparation Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Medium Web Scraper & Search API Assignment', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Read Markdown content
    if not os.path.exists('QUIZ_PREP.md'):
        print("❌ QUIZ_PREP.md not found")
        return

    with open('QUIZ_PREP.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Already handled main title, maybe skip or add as H1
            pass 
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('|'):
            # Simple table handling (skip for now, just add as text or handle simpler)
            # Actually, let's just add it as pre-formatted text if it looks like a table
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
        elif line.startswith('**Q:'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 50, 150) # Dark Blue
        elif line.startswith('> **A:'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(line.replace('> ', ''))
            run.italic = True
        elif line.startswith('```'):
            # Code block marker, ignore or switch style
            pass
        elif line.startswith('    ') or line.startswith('\t'):
             # Code lines
             p = doc.add_paragraph(line)
             p.style = 'No Spacing'
             for run in p.runs:
                 run.font.name = 'Courier New'
                 run.font.size = Pt(10)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
             doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save('Quiz_Preparation_Guide.docx')
    print("✅ Created: Quiz_Preparation_Guide.docx")

if __name__ == '__main__':
    create_prep_guide()
